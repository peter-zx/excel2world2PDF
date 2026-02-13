"""
端到端测试：模拟完整用户流程
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from io import BytesIO
from docx import Document
import pandas as pd

from src.services.template_service import template_service
from src.services.excel_service import excel_service
from src.services.word_service import word_service


def create_test_contract():
    """创建测试合同"""
    doc = Document()
    doc.add_paragraph("劳动合同")
    doc.add_paragraph("乙    方： 陈长  身份证号： 420115197806100095")
    doc.add_paragraph("自  2025  年   7  月  1  日起至  2028  年   6  月  30  日止")
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def create_test_excel():
    """创建测试Excel"""
    data = {
        "姓名": ["张三", "李四"],
        "身份证号": ["110101199001011234", "110101199002022345"],
        "起始年": ["2024", "2024"],
    }
    df = pd.DataFrame(data)
    
    output = BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    return output.getvalue()


def test_e2e():
    """端到端测试"""
    print("=" * 60)
    print("端到端测试：模拟完整用户流程")
    print("=" * 60)
    
    # Step 1: 创建模板
    print("\n[Step 1] 创建测试合同...")
    contract_bytes = create_test_contract()
    
    # 查看原文
    doc = Document(BytesIO(contract_bytes))
    para1_text = doc.paragraphs[1].text
    para2_text = doc.paragraphs[2].text
    print(f"  段落1: {para1_text}")
    print(f"  段落2: {para2_text}")
    
    # Step 2: 配置映射
    print("\n[Step 2] 配置位置映射...")
    
    # 精确查找位置
    name_pos = para1_text.find("陈长")
    id_pos = para1_text.find("420115197806100095")
    year_pos = para2_text.find("2025")
    
    location_mapping = {
        "姓名": {
            "element_id": "para_1",
            "start": name_pos,
            "end": name_pos + 2,
            "length": 2,
            "original_text": "陈长"
        },
        "身份证号": {
            "element_id": "para_1",
            "start": id_pos,
            "end": id_pos + 18,
            "length": 18,
            "original_text": "420115197806100095"
        },
        "起始年": {
            "element_id": "para_2",
            "start": year_pos,
            "end": year_pos + 4,
            "length": 4,
            "original_text": "2025"
        }
    }
    
    print(f"  位置映射:")
    for k, v in location_mapping.items():
        print(f"    {k}: {v['element_id']} [{v['start']}:{v['end']}] = '{v['original_text']}'")
    
    # Step 3: 保存模板
    print("\n[Step 3] 保存模板...")
    config = template_service.create_location_template(
        template_name="测试合同模板",
        original_filename="test_contract.docx",
        docx_bytes=contract_bytes,
        location_mapping=location_mapping,
        description="测试用"
    )
    print(f"  模板ID: {config.template_id}")
    print(f"  模板名称: {config.template_name}")
    
    # Step 4: 读取模板
    print("\n[Step 4] 验证模板读取...")
    loaded_config = template_service.load_config(config.template_id)
    print(f"  加载的映射数量: {len(loaded_config.location_mapping)}")
    for k, v in loaded_config.location_mapping.items():
        print(f"    {k}: {v}")
    
    # Step 5: 创建Excel数据
    print("\n[Step 5] 创建Excel数据...")
    excel_bytes = create_test_excel()
    df = pd.read_excel(BytesIO(excel_bytes))
    print(f"  数据:\n{df}")
    
    # Step 6: 生成合同
    print("\n[Step 6] 生成合同...")
    template_bytes = template_service.get_template_bytes(config.template_id)
    data_list = excel_service.dataframe_to_dict_list(df)
    
    print(f"  数据列表: {data_list}")
    
    files = word_service.batch_generate_by_location(
        template_bytes,
        data_list,
        loaded_config.location_mapping
    )
    
    print(f"\n  生成结果: {len(files)} 份合同")
    
    # 验证结果
    for i, (filename, file_bytes) in enumerate(files):
        result_doc = Document(BytesIO(file_bytes))
        print(f"\n  [{i+1}] {filename}:")
        for j, para in enumerate(result_doc.paragraphs):
            print(f"      段落{j}: {para.text}")
    
    # 清理
    template_service.delete_template(config.template_id)
    print("\n[清理] 测试模板已删除")
    
    print("\n" + "=" * 60)
    print("测试完成!")
    print("=" * 60)


if __name__ == "__main__":
    test_e2e()
