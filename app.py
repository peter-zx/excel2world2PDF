"""
合同自动填写工具 - 优化版
- A4预览布局
- 卡片式段落展示 + 引导线
- 双列布局：预览 + 配置
"""
import streamlit as st
import pandas as pd
import zipfile
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Optional
import re

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from src.services.template_service import template_service
from src.services.excel_service import excel_service
from src.services.word_service import word_service
from src.models.schemas import TemplateConfig


# ==================== 页面配置 ====================
st.set_page_config(
    page_title="合同自动填写工具",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS样式
CUSTOM_CSS = """
<style>
/* A4容器样式 */
.a4-container {
    background: white;
    width: 100%;
    max-width: 800px;
    min-height: 600px;
    padding: 40px 50px;
    margin: 0 auto;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    border: 1px solid #e0e0e0;
    font-family: "SimSun", "宋体", serif;
}

/* 段落卡片样式 */
.para-card {
    padding: 12px 16px;
    margin: 8px 0;
    border-radius: 8px;
    border: 1px solid #e8e8e8;
    background: #fafafa;
    transition: all 0.2s;
    position: relative;
}

.para-card:hover {
    border-color: #1890ff;
    background: #f0f7ff;
}

.para-card.selected {
    border-color: #1890ff;
    background: #e6f4ff;
    box-shadow: 0 0 0 2px rgba(24,144,255,0.2);
}

/* 引导线 */
.connector {
    position: absolute;
    right: -20px;
    top: 50%;
    width: 20px;
    height: 2px;
    background: #1890ff;
}

.connector::after {
    content: "→";
    position: absolute;
    right: -8px;
    top: -10px;
    color: #1890ff;
    font-size: 16px;
}

/* 段落序号 */
.para-num {
    display: inline-block;
    width: 28px;
    height: 28px;
    line-height: 28px;
    text-align: center;
    background: #1890ff;
    color: white;
    border-radius: 50%;
    font-size: 12px;
    font-weight: bold;
    margin-right: 12px;
}

/* 配置面板样式 */
.config-panel {
    background: #f8f9fa;
    padding: 20px;
    border-radius: 8px;
    border: 1px solid #e8e8e8;
}

/* 检测项样式 */
.detect-item {
    padding: 8px 12px;
    margin: 6px 0;
    background: white;
    border-radius: 6px;
    border: 1px solid #e8e8e8;
    display: flex;
    align-items: center;
    justify-content: space-between;
}

.detect-text {
    font-family: monospace;
    background: #fff3cd;
    padding: 2px 6px;
    border-radius: 3px;
}

/* 已映射标签 */
.mapped-tag {
    display: inline-block;
    background: #52c41a;
    color: white;
    padding: 2px 8px;
    border-radius: 4px;
    font-size: 12px;
    margin-left: 8px;
}
</style>
"""


# ==================== Session State ====================
def init_session_state():
    defaults = {
        "current_step": "template",
        "uploaded_template_bytes": None,
        "template_name": "",
        "description": "",
        "doc_elements": [],
        "location_mapping": {},
        "selected_element_id": None,
        "uploaded_df": None,
        "selected_template": None,
        "generated_files": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_session_state()


# ==================== 文档解析 ====================
def parse_doc_elements(file_bytes: bytes) -> List[Dict]:
    """解析Word文档，返回元素列表"""
    from docx import Document
    
    doc = Document(BytesIO(file_bytes))
    elements = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            elements.append({
                "type": "paragraph",
                "index": para_idx,
                "element_id": f"para_{para_idx}",
                "text": para.text,
            })
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    elements.append({
                        "type": "table_cell",
                        "index": f"t{table_idx}_r{row_idx}_c{cell_idx}",
                        "element_id": f"cell_{table_idx}_{row_idx}_{cell_idx}",
                        "text": cell.text,
                    })
    
    return elements


def extract_candidates(text: str) -> List[Dict]:
    """从文本中提取候选替换内容"""
    candidates = []
    seen = set()
    
    patterns = [
        (r'\d{15,18}[Xx]?', "身份证号"),
        (r'1[3-9]\d{9}', "手机号"),
        (r'20\d{2}', "年份"),
        (r'\d{4,}(?:\.\d{1,2})?', "金额/数字"),
    ]
    
    for pattern, ptype in patterns:
        for m in re.finditer(pattern, text):
            val = m.group()
            if val not in seen:
                candidates.append({
                    "text": val,
                    "type": ptype,
                    "start": m.start(),
                    "end": m.end()
                })
                seen.add(val)
    
    return candidates


def generate_excel_template(mapping: Dict[str, str]) -> bytes:
    """生成Excel模板"""
    df = pd.DataFrame(columns=list(mapping.keys()))
    example = {k: v if v.isdigit() else f"示例{k}" for k, v in mapping.items()}
    df = pd.concat([df, pd.DataFrame([example])], ignore_index=True)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='数据')
    output.seek(0)
    return output.getvalue()


# ==================== 侧边栏 ====================
def render_sidebar():
    st.sidebar.title("📋 导航")
    
    steps = {
        "template": "📄 模板管理",
        "data": "📊 数据导入",
        "generate": "🚀 批量生成"
    }
    
    for step_key, step_name in steps.items():
        btn_type = "primary" if st.session_state.current_step == step_key else "secondary"
        if st.sidebar.button(step_name, key=f"nav_{step_key}", use_container_width=True, type=btn_type):
            st.session_state.current_step = step_key
            st.rerun()
    
    st.sidebar.divider()
    templates = template_service.list_templates()
    st.sidebar.info(f"📚 已保存: {len(templates)} 个模板")


# ==================== 模板管理页面 ====================
def render_template_page():
    st.header("📄 步骤1: 模板管理")
    
    # 已保存模板
    with st.expander("📚 已保存模板", expanded=False):
        templates = template_service.list_templates()
        if templates:
            for tpl in templates:
                c1, c2, c3 = st.columns([3, 2, 1])
                c1.write(f"**{tpl.template_name}**")
                c2.write(f"{len(tpl.location_mapping)} 个变量")
                if c3.button("🗑️", key=f"del_{tpl.template_id}"):
                    template_service.delete_template(tpl.template_id)
                    st.rerun()
        else:
            st.info("暂无保存的模板")
    
    st.divider()
    
    # 上传
    uploaded_file = st.file_uploader("📤 上传Word合同 (.docx)", type=["docx"])
    
    if uploaded_file:
        file_bytes = uploaded_file.getvalue()
        
        if st.session_state.uploaded_template_bytes != file_bytes:
            st.session_state.uploaded_template_bytes = file_bytes
            st.session_state.doc_elements = parse_doc_elements(file_bytes)
            st.session_state.location_mapping = {}
            st.session_state.template_name = Path(uploaded_file.name).stem
            st.session_state.selected_element_id = None
        
        st.markdown(CUSTOM_CSS, unsafe_allow_html=True)
        
        # ========== 双列布局 ==========
        col_preview, col_config = st.columns([3, 2])
        
        # ===== 左侧：A4预览 =====
        with col_preview:
            st.subheader("📄 合同预览")
            st.caption("点击段落编号选择，在右侧配置映射")
            
            with st.container():
                for i, elem in enumerate(st.session_state.doc_elements):
                    elem_id = elem["element_id"]
                    text = elem["text"]
                    is_selected = st.session_state.selected_element_id == elem_id
                    
                    # 检查是否已映射
                    mapped_vars = [
                        k for k, v in st.session_state.location_mapping.items() 
                        if v["element_id"] == elem_id
                    ]
                    
                    # 段落卡片
                    c1, c2 = st.columns([0.5, 9.5])
                    
                    with c1:
                        # 序号按钮
                        btn_label = f"{'📍' if is_selected else str(i+1)}"
                        if st.button(btn_label, key=f"sel_{elem_id}"):
                            st.session_state.selected_element_id = elem_id
                            st.rerun()
                    
                    with c2:
                        # 内容显示
                        display_text = text[:80] + "..." if len(text) > 80 else text
                        
                        if is_selected:
                            st.markdown(f"**{display_text}**")
                        else:
                            st.text(display_text)
                        
                        if mapped_vars:
                            st.markdown(" ".join([f"`{v}`" for v in mapped_vars]))
                    
                    st.divider()
        
        # ===== 右侧：配置面板 =====
        with col_config:
            st.subheader("🏷️ 配置映射")
            
            if st.session_state.selected_element_id:
                elem_id = st.session_state.selected_element_id
                elem = next((e for e in st.session_state.doc_elements if e["element_id"] == elem_id), None)
                
                if elem:
                    st.markdown(f"**选中段落:**")
                    st.info(elem["text"][:100] + ("..." if len(elem["text"]) > 100 else ""))
                    
                    # === 自定义映射 ===
                    st.markdown("#### ✏️ 自定义映射")
                    st.caption("手动输入要替换的文本和变量名")
                    
                    c1, c2 = st.columns(2)
                    with c1:
                        custom_text = st.text_input("原文本", key="custom_text", placeholder="如：陈长")
                    with c2:
                        custom_var = st.text_input("变量名", key="custom_var", placeholder="如：姓名")
                    
                    if st.button("➕ 添加自定义映射", type="primary", use_container_width=True):
                        if custom_text and custom_var:
                            pos = elem["text"].find(custom_text)
                            if pos >= 0:
                                st.session_state.location_mapping[custom_var] = {
                                    "element_id": elem_id,
                                    "start": pos,
                                    "end": pos + len(custom_text),
                                    "length": len(custom_text),
                                    "original_text": custom_text
                                }
                                st.success(f"✅ 已添加: {custom_var} = {custom_text}")
                                st.rerun()
                            else:
                                st.error(f"未找到文本: {custom_text}")
                    
                    st.divider()
                    
                    # === 智能检测 ===
                    st.markdown("#### 🔍 智能检测")
                    st.caption("自动识别可替换内容")
                    
                    candidates = extract_candidates(elem["text"])
                    
                    if candidates:
                        for cand in candidates:
                            c1, c2, c3 = st.columns([3, 2, 1])
                            with c1:
                                st.markdown(f"`{cand['text']}` ({cand['type']})")
                            with c2:
                                var_input = st.text_input(
                                    "变量名", 
                                    key=f"det_var_{cand['start']}",
                                    placeholder="输入变量名",
                                    label_visibility="collapsed"
                                )
                            with c3:
                                if st.button("使用", key=f"det_add_{cand['start']}"):
                                    if var_input:
                                        st.session_state.location_mapping[var_input] = {
                                            "element_id": elem_id,
                                            "start": cand["start"],
                                            "end": cand["end"],
                                            "length": cand["end"] - cand["start"],
                                            "original_text": cand["text"]
                                        }
                                        st.rerun()
                    else:
                        st.info("未检测到可替换内容")
            else:
                st.info("👈 请在左侧点击段落编号选择")
            
            # === 已配置映射 ===
            st.divider()
            st.markdown("#### 📋 已配置映射")
            
            if st.session_state.location_mapping:
                for var_name, loc in st.session_state.location_mapping.items():
                    c1, c2 = st.columns([4, 1])
                    with c1:
                        st.write(f"**{var_name}** = `{loc['original_text']}`")
                    with c2:
                        if st.button("🗑️", key=f"del_map_{var_name}"):
                            del st.session_state.location_mapping[var_name]
                            st.rerun()
                
                # 下载模板
                st.divider()
                if st.button("📥 下载Excel模板", use_container_width=True):
                    simple_map = {k: v["original_text"] for k, v in st.session_state.location_mapping.items()}
                    excel_bytes = generate_excel_template(simple_map)
                    st.download_button(
                        label="📥 点击下载",
                        data=excel_bytes,
                        file_name=f"{st.session_state.template_name}_模板.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            else:
                st.info("暂无映射配置")
        
        # === 保存模板 ===
        st.divider()
        st.subheader("💾 保存模板")
        
        c1, c2 = st.columns([2, 1])
        with c1:
            st.session_state.template_name = st.text_input("模板名称 *", value=st.session_state.template_name)
        with c2:
            st.session_state.description = st.text_input("描述", value=st.session_state.description)
        
        if st.button("💾 保存模板", type="primary", use_container_width=True):
            if not st.session_state.template_name:
                st.error("请输入模板名称")
            elif not st.session_state.location_mapping:
                st.error("请至少添加一个映射")
            else:
                try:
                    config = template_service.create_location_template(
                        template_name=st.session_state.template_name,
                        original_filename=uploaded_file.name,
                        docx_bytes=file_bytes,
                        location_mapping=st.session_state.location_mapping,
                        description=st.session_state.description
                    )
                    st.success(f"✅ 保存成功！ID: {config.template_id}")
                except Exception as e:
                    st.error(f"保存失败: {e}")


# ==================== 数据导入页面 ====================
def render_data_page():
    st.header("📊 步骤2: 数据导入")
    
    templates = template_service.list_templates()
    if not templates:
        st.warning("⚠️ 请先创建模板")
        return
    
    st.subheader("📚 选择模板")
    template_options = {f"{t.template_name}": t for t in templates}
    selected_key = st.selectbox("选择模板", options=list(template_options.keys()))
    
    selected = template_options[selected_key]
    st.session_state.selected_template = selected
    
    cols = list(selected.location_mapping.keys())
    st.info(f"**需要列:** {', '.join(cols)}")
    
    if st.button("📥 下载Excel模板"):
        simple_map = {k: v.get("original_text", "") for k, v in selected.location_mapping.items()}
        excel_bytes = generate_excel_template(simple_map)
        st.download_button(
            label="📥 下载",
            data=excel_bytes,
            file_name=f"{selected.template_name}_模板.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    
    st.divider()
    
    st.subheader("📤 上传Excel")
    excel_file = st.file_uploader("选择Excel文件", type=["xlsx", "xls"])
    
    if excel_file:
        df, error = excel_service.read_excel(excel_file.getvalue(), excel_file.name)
        if error:
            st.error(f"读取失败: {error}")
            return
        
        st.session_state.uploaded_df = df
        st.dataframe(df, use_container_width=True)
        st.info(f"共 {len(df)} 条记录")


# ==================== 批量生成页面 ====================
def render_generate_page():
    st.header("🚀 步骤3: 批量生成")
    
    if not st.session_state.selected_template:
        st.warning("⚠️ 请先选择模板")
        return
    
    if st.session_state.uploaded_df is None:
        st.warning("⚠️ 请先上传数据")
        return
    
    template = st.session_state.selected_template
    df = st.session_state.uploaded_df
    
    c1, c2 = st.columns(2)
    c1.info(f"**模板:** {template.template_name}")
    c2.info(f"**数据:** {len(df)} 条")
    
    if st.button("🚀 开始生成", type="primary", use_container_width=True):
        with st.spinner("生成中..."):
            try:
                template_bytes = template_service.get_template_bytes(template.template_id)
                if not template_bytes:
                    st.error("模板不存在")
                    return
                
                data_list = excel_service.dataframe_to_dict_list(df)
                files = word_service.batch_generate_by_location(
                    template_bytes, data_list, template.location_mapping
                )
                
                st.session_state.generated_files = files
                st.success(f"✅ 成功 {len(files)} 份")
                
            except Exception as e:
                st.error(f"失败: {e}")
    
    if st.session_state.generated_files:
        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as z:
            for fn, fb in st.session_state.generated_files:
                z.writestr(fn, fb)
        zip_buf.seek(0)
        
        st.download_button(
            label="📦 下载全部",
            data=zip_buf,
            file_name=f"合同_{datetime.now():%Y%m%d_%H%M%S}.zip",
            mime="application/zip",
            use_container_width=True
        )


# ==================== 主函数 ====================
def main():
    st.title("📝 合同自动填写工具")
    st.markdown("上传合同 → 选择段落 → 配置映射 → 批量生成")
    
    render_sidebar()
    st.divider()
    
    if st.session_state.current_step == "template":
        render_template_page()
    elif st.session_state.current_step == "data":
        render_data_page()
    elif st.session_state.current_step == "generate":
        render_generate_page()


if __name__ == "__main__":
    main()
