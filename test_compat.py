"""
测试旧格式模板兼容性
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from io import BytesIO
from docx import Document
import pandas as pd

from src.services.template_service import template_service
from src.services.excel_service import excel_service
from src.services.word_service import word_service
from src.models.schemas import TemplateConfig


def create_test_contract():
    """创建测试合同"""
    doc = Document()
    doc.add_paragraph("劳动合同")
    doc.add_paragraph("乙    方： 陈长  身份证号： 420115197806100095")
    doc.add_paragraph("自  2025  年   7  月  1  日起至  2028  年   6  月  30  日止")
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def test_text_mapping_compat():
    """测试旧格式（text_mapping）兼容性"""
    print("=" * 60)
    print("测试旧格式模板兼容性")
    print("=" * 60)
    
    # 创建合同
    contract_bytes = create_test_contract()
    
    # 使用旧格式：text_mapping
    text_mapping = {
        "姓名": "陈长",
        "身份证号": "420115197806100095",
        "起始年": "2025"
    }
    
    # 手动创建旧格式配置（模拟用户之前保存的模板）
    import json
    import uuid
    from src.config import TEMPLATES_DIR, CONFIGS_DIR
    
    template_id = str(uuid.uuid4())[:8]
    template_filename = f"{template_id}_test.docx"
    
    # 保存模板文件
    template_path = TEMPLATES_DIR / template_filename
    with open(template_path, "wb") as f:
        f.write(contract_bytes)
    
    # 保存旧格式配置
    config_data = {
        "template_id": template_id,
        "template_name": "旧格式测试模板",
        "original_filename": "test.docx",
        "template_filename": template_filename,
        "location_mapping": {},  # 空
        "text_mapping": text_mapping,  # 旧格式
        "created_at": "2024-01-01",
        "updated_at": "2024-01-01",
        "description": "旧格式测试"
    }
    
    config_path = CONFIGS_DIR / f"{template_id}.json"
    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n[Step 1] 创建旧格式模板: {template_id}")
    print(f"  text_mapping: {text_mapping}")
    
    # 加载模板
    template = template_service.load_config(template_id)
    mapping_info = template.get_mapping()
    print(f"\n[Step 2] 加载模板")
    print(f"  映射类型: {mapping_info['type']}")
    
    # 创建数据
    data_list = [
        {"姓名": "张三", "身份证号": "110101199001011234", "起始年": "2024"},
        {"姓名": "李四", "身份证号": "110101199002022345", "起始年": "2024"},
    ]
    
    print(f"\n[Step 3] 生成合同")
    
    # 获取模板文件
    template_bytes = template_service.get_template_bytes(template_id)
    
    # 根据映射类型生成
    if mapping_info['type'] == 'text':
        files = word_service.batch_generate_by_text(
            template_bytes, data_list, mapping_info['data']
        )
    else:
        files = word_service.batch_generate_by_location(
            template_bytes, data_list, mapping_info['data']
        )
    
    print(f"  生成结果: {len(files)} 份")
    
    # 验证
    all_passed = True
    for i, (filename, file_bytes) in enumerate(files):
        result_doc = Document(BytesIO(file_bytes))
        para1 = result_doc.paragraphs[1].text
        para2 = result_doc.paragraphs[2].text
        
        print(f"\n  [{i+1}] {filename}:")
        print(f"      段落1: {para1}")
        print(f"      段落2: {para2}")
        
        expected_name = data_list[i]["姓名"]
        expected_id = data_list[i]["身份证号"]
        expected_year = data_list[i]["起始年"]
        
        if expected_name not in para1:
            print(f"      [FAIL] 姓名未替换")
            all_passed = False
        if expected_id not in para1:
            print(f"      [FAIL] 身份证未替换")
            all_passed = False
        if expected_year not in para2:
            print(f"      [FAIL] 年份未替换")
            all_passed = False
    
    # 清理
    template_service.delete_template(template_id)
    
    if all_passed:
        print("\n>>> 旧格式兼容性测试通过！")
    else:
        print("\n>>> 旧格式兼容性测试失败！")
    
    print("=" * 60)


if __name__ == "__main__":
    test_text_mapping_compat()
