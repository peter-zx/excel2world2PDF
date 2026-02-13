"""
测试格式保留
验证字体、字号、下划线等格式是否保留
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

from src.services.word_service import word_service


def create_formatted_contract():
    """创建带格式的合同模板"""
    doc = Document()
    
    # 添加标题（大字体、加粗）
    title = doc.add_paragraph()
    run = title.add_run("劳动合同")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    
    # 添加带下划线的段落
    para1 = doc.add_paragraph()
    
    # "乙    方：" - 普通格式
    run1 = para1.add_run("乙    方：")
    run1.font.name = "宋体"
    run1.font.size = Pt(12)
    
    # "陈长" - 带下划线
    run2 = para1.add_run("陈长")
    run2.font.name = "宋体"
    run2.font.size = Pt(12)
    run2.font.underline = WD_UNDERLINE.SINGLE
    
    # "  身份证号：" - 普通格式
    run3 = para1.add_run("  身份证号：")
    run3.font.name = "宋体"
    run3.font.size = Pt(12)
    
    # 身份证号 - 带下划线
    run4 = para1.add_run("420115197806100095")
    run4.font.name = "宋体"
    run4.font.size = Pt(12)
    run4.font.underline = WD_UNDERLINE.SINGLE
    
    # 添加日期段落
    para2 = doc.add_paragraph()
    run5 = para2.add_run("自  ")
    run5.font.name = "宋体"
    run5.font.size = Pt(12)
    
    run6 = para2.add_run("2025")
    run6.font.name = "宋体"
    run6.font.size = Pt(12)
    run6.font.underline = WD_UNDERLINE.SINGLE
    
    run7 = para2.add_run("  年  ")
    run7.font.name = "宋体"
    run7.font.size = Pt(12)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def check_format(paragraph, idx):
    """检查段落的格式"""
    print(f"\n段落{idx}的run信息:")
    for i, run in enumerate(paragraph.runs):
        print(f"  Run {i}: '{run.text}'")
        print(f"    字体: {run.font.name}")
        print(f"    字号: {run.font.size.pt if run.font.size else 'None'}pt")
        print(f"    加粗: {run.font.bold}")
        print(f"    下划线: {run.font.underline}")
        print(f"    斜体: {run.font.italic}")


def test_format_preservation():
    """测试格式保留"""
    print("=" * 60)
    print("测试格式保留")
    print("=" * 60)
    
    # 创建带格式的模板
    template_bytes = create_formatted_contract()
    
    # 查看原始格式
    print("\n[原始模板格式]")
    doc = Document(BytesIO(template_bytes))
    for i, para in enumerate(doc.paragraphs):
        check_format(para, i)
    
    # 定义替换
    text_mapping = {
        "姓名": "陈长",
        "身份证号": "420115197806100095",
        "起始年": "2025"
    }
    
    data = {
        "姓名": "张三",
        "身份证号": "110101199001011234",
        "起始年": "2024"
    }
    
    # 执行替换
    print("\n[执行替换]")
    print(f"  姓名: 陈长 -> 张三")
    print(f"  身份证号: 420115197806100095 -> 110101199001011234")
    print(f"  起始年: 2025 -> 2024")
    
    result_bytes = word_service.replace_preserving_format(
        template_bytes, data, text_mapping=text_mapping
    )
    
    # 查看替换后的格式
    print("\n[替换后格式]")
    result_doc = Document(BytesIO(result_bytes))
    for i, para in enumerate(result_doc.paragraphs):
        check_format(para, i)
    
    # 验证格式保留
    print("\n[格式验证]")
    
    # 检查姓名是否保留下划线
    name_run = None
    for run in result_doc.paragraphs[1].runs:
        if "张三" in run.text:
            name_run = run
            break
    
    if name_run and name_run.font.underline:
        print("  姓名: 下划线保留 [OK]")
    else:
        print("  姓名: 下划线丢失 [FAIL]")
    
    # 检查身份证是否保留下划线
    id_run = None
    for run in result_doc.paragraphs[1].runs:
        if "110101" in run.text:
            id_run = run
            break
    
    if id_run and id_run.font.underline:
        print("  身份证号: 下划线保留 [OK]")
    else:
        print("  身份证号: 下划线丢失 [FAIL]")
    
    # 检查字号
    all_12pt = True
    for para in result_doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size.pt != 12:
                if run.text.strip():  # 忽略空白run
                    all_12pt = False
    
    if all_12pt:
        print("  字号: 全部12pt [OK]")
    else:
        print("  字号: 有变化 [FAIL]")
    
    # 保存结果
    with open("test_format_result.docx", "wb") as f:
        f.write(result_bytes)
    print("\n结果已保存到 test_format_result.docx")


if __name__ == "__main__":
    test_format_preservation()
