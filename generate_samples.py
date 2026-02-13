"""
生成示例模板和Excel数据文件
运行此脚本创建测试用的示例文件
"""
import sys
from pathlib import Path

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent))

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_sample_template():
    """创建示例Word合同模板（包含【】占位符）"""
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('劳动合同', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 合同编号
    p0 = doc.add_paragraph()
    p0.add_run('合同编号：【合同编号】')
    
    doc.add_paragraph()
    
    # 甲方信息
    p1 = doc.add_paragraph()
    p1.add_run('甲    方：【甲方名称】')
    
    # 乙方信息
    p2 = doc.add_paragraph()
    p2.add_run('乙    方：【姓名】________________')
    
    p3 = doc.add_paragraph()
    p3.add_run('身份证号：【身份证号】________________')
    
    p4 = doc.add_paragraph()
    p4.add_run('联系方式：【联系电话】________________')
    
    p5 = doc.add_paragraph()
    p5.add_run('家庭住址：【地址】________________')
    
    doc.add_paragraph()
    
    # 合同期限
    doc.add_heading('一、合同期限', level=1)
    p6 = doc.add_paragraph()
    p6.add_run('甲乙双方选择合同期限为：')
    
    p7 = doc.add_paragraph()
    p7.add_run('自【起始年】年【起始月】月【起始日】日起至')
    p7.add_run('【结束年】年【结束月】月【结束日】日止，共【合同年限】年。')
    
    doc.add_heading('二、工作内容', level=1)
    p8 = doc.add_paragraph()
    p8.add_run('工作岗位为【职位】')
    
    doc.add_heading('三、劳动报酬', level=1)
    p9 = doc.add_paragraph()
    p9.add_run('工资（税前）【月薪】元/月')
    
    doc.add_heading('四、其他约定', level=1)
    p10 = doc.add_paragraph('【备注】')
    
    doc.add_paragraph()
    
    # 签名区
    doc.add_paragraph('甲方（盖章）：____________________')
    p_sign = doc.add_paragraph()
    p_sign.add_run('乙方（签字）：【姓名】________________')
    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.add_run('签订日期：【签订日期】')
    
    # 保存文件
    output_dir = Path(__file__).parent / "samples"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "合同模板示例.docx"
    doc.save(output_path)
    print(f"[OK] 已创建示例模板：{output_path}")
    return output_path


def create_sample_excel():
    """创建示例Excel数据文件"""
    data = {
        '合同编号': ['HT-2024-001', 'HT-2024-002', 'HT-2024-003'],
        '甲方名称': ['XX科技有限公司', 'XX科技有限公司', 'XX科技有限公司'],
        '姓名': ['张三', '李四', '王五'],
        '身份证号': ['110101199001011234', '110101199002022345', '110101199003033456'],
        '联系电话': ['13800138001', '13800138002', '13800138003'],
        '地址': ['北京市朝阳区XX路1号', '北京市海淀区XX路2号', '北京市西城区XX路3号'],
        '起始年': ['2024', '2024', '2024'],
        '起始月': ['01', '02', '03'],
        '起始日': ['01', '01', '01'],
        '结束年': ['2027', '2027', '2027'],
        '结束月': ['01', '02', '03'],
        '结束日': ['01', '01', '01'],
        '合同年限': ['3', '3', '3'],
        '职位': ['工程师', '经理', '会计'],
        '月薪': ['15000', '18000', '12000'],
        '签订日期': ['2024-01-01', '2024-02-01', '2024-03-01'],
        '备注': ['无', '优秀员工', '无']
    }
    
    df = pd.DataFrame(data)
    output_dir = Path(__file__).parent / "samples"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "员工数据示例.xlsx"
    df.to_excel(output_path, index=False)
    print(f"[OK] 已创建示例数据：{output_path}")
    return output_path


if __name__ == '__main__':
    print("=" * 50)
    print("正在生成示例文件...")
    print("=" * 50)
    
    template_path = create_sample_template()
    excel_path = create_sample_excel()
    
    print("\n" + "=" * 50)
    print("示例文件生成完成！")
    print("=" * 50)
    print(f"\n模板文件: {template_path}")
    print(f"数据文件: {excel_path}")
    print("\n使用方法:")
    print("1. 运行: streamlit run app.py")
    print("2. 在WebUI中上传生成的模板和数据文件")
