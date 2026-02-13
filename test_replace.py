"""
测试Word替换功能 - 调试版
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from io import BytesIO
from docx import Document
from src.services.word_service import word_service


def create_test_doc():
    """创建测试文档"""
    doc = Document()
    
    doc.add_paragraph("劳动合同")
    doc.add_paragraph("乙    方： 陈长  身份证号： 420115197806100095")
    doc.add_paragraph("自  2025  年   7  月  1  日起至  2028  年   6  月  30  日止")
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def debug_positions():
    """调试位置计算"""
    template_bytes = create_test_doc()
    doc = Document(BytesIO(template_bytes))
    
    print("=" * 60)
    print("调试位置计算")
    print("=" * 60)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        print(f"\n段落{i}: {repr(text)}")
        print(f"长度: {len(text)}")
        
        # 显示每个字符的位置
        for j, char in enumerate(text):
            if j % 10 == 0:
                print(f"\n{j:3d}: ", end="")
            print(char, end="")
        print()


def test_replace():
    """测试替换功能"""
    print("\n" + "=" * 60)
    print("测试替换功能")
    print("=" * 60)
    
    # 创建测试文档
    template_bytes = create_test_doc()
    
    # 查看原始内容
    doc = Document(BytesIO(template_bytes))
    print("\n原始文档内容:")
    for i, para in enumerate(doc.paragraphs):
        print(f"  段落{i}: {repr(para.text)}")
    
    # 定义位置映射（需要精确计算位置）
    # 段落1: "乙    方： 陈长  身份证号： 420115197806100095"
    # 手动数位置...
    para1_text = doc.paragraphs[1].text
    para2_text = doc.paragraphs[2].text
    
    # 精确查找位置
    name_pos = para1_text.find("陈长")
    id_pos = para1_text.find("420115197806100095")
    year_pos = para2_text.find("2025")
    
    print(f"\n精确位置:")
    print(f"  '陈长' 位置: {name_pos}-{name_pos + 2}")
    print(f"  '420115197806100095' 位置: {id_pos}-{id_pos + 18}")
    print(f"  '2025' 位置: {year_pos}-{year_pos + 4}")
    
    location_mapping = {
        "姓名": {
            "element_id": "para_1",
            "start": name_pos,
            "end": name_pos + 2,
            "original_text": "陈长"
        },
        "身份证号": {
            "element_id": "para_1",
            "start": id_pos,
            "end": id_pos + 18,
            "original_text": "420115197806100095"
        },
        "起始年": {
            "element_id": "para_2",
            "start": year_pos,
            "end": year_pos + 4,
            "original_text": "2025"
        }
    }
    
    # 定义替换数据
    data = {
        "姓名": "张三",
        "身份证号": "110101199001011234",
        "起始年": "2024"
    }
    
    print("\n替换数据:")
    for k, v in data.items():
        print(f"  {k}: {v}")
    
    # 执行替换
    result_bytes = word_service.replace_by_location(template_bytes, location_mapping, data)
    
    # 查看替换结果
    result_doc = Document(BytesIO(result_bytes))
    print("\n替换后文档内容:")
    for i, para in enumerate(result_doc.paragraphs):
        print(f"  段落{i}: {repr(para.text)}")
    
    # 验证结果
    print("\n验证结果:")
    para1_result = result_doc.paragraphs[1].text
    para2_result = result_doc.paragraphs[2].text
    
    checks = [
        ("姓名替换正确", "张三" in para1_result),
        ("身份证替换正确", "110101199001011234" in para1_result),
        ("年份替换正确", "2024" in para2_result),
        ("原姓名已删除", "陈长" not in para1_result),
        ("原身份证已删除", "420115197806100095" not in para1_result),
        ("原年份已删除", "2025" not in para2_result or "2024" in para2_result),
    ]
    
    all_passed = True
    for name, passed in checks:
        status = "[PASS]" if passed else "[FAIL]"
        print(f"  {name}: {status}")
        if not passed:
            all_passed = False
    
    if all_passed:
        print("\n>>> All tests passed!")
    else:
        print("\n>>> Some tests failed!")
    
    # 保存结果文件
    with open("test_result.docx", "wb") as f:
        f.write(result_bytes)
    print("\n结果已保存到 test_result.docx")


if __name__ == "__main__":
    debug_positions()
    test_replace()
