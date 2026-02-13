"""
模板管理页面
"""
import streamlit as st
from pathlib import Path
from io import BytesIO
from docx import Document
from typing import Dict, List

from src.services.template_service import template_service
from src.utils import extract_candidates, generate_excel_template
from src.components import show_success, show_error, show_warning, show_info


# 页面样式
PAGE_STYLE = """
<style>
.para-card {
    padding: 12px 16px;
    margin: 8px 0;
    border-radius: 8px;
    border: 1px solid #e8e8e8;
    background: #fafafa;
}
.para-card:hover {
    border-color: #1890ff;
    background: #f0f7ff;
}
.para-card.selected {
    border-color: #1890ff;
    background: #e6f4ff;
}
</style>
"""


def parse_doc_elements(file_bytes: bytes) -> List[Dict]:
    """解析Word文档，返回元素列表"""
    doc = Document(BytesIO(file_bytes))
    elements = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            elements.append({
                "type": "paragraph",
                "index": para_idx,
                "element_id": f"para_{para_idx}",
                "text": para.text,
            })
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    elements.append({
                        "type": "table_cell",
                        "index": f"t{table_idx}_r{row_idx}_c{cell_idx}",
                        "element_id": f"cell_{table_idx}_{row_idx}_{cell_idx}",
                        "text": cell.text,
                    })
    
    return elements


def render_saved_templates():
    """渲染已保存模板列表"""
    with st.expander("📚 已保存模板", expanded=False):
        templates = template_service.list_templates()
        if templates:
            for tpl in templates:
                c1, c2, c3 = st.columns([3, 2, 1])
                c1.write(f"**{tpl.template_name}**")
                mapping_info = tpl.get_mapping()
                c2.write(f"{len(mapping_info['data'])} 个变量")
                if c3.button("🗑️", key=f"del_{tpl.template_id}"):
                    template_service.delete_template(tpl.template_id)
                    st.rerun()
        else:
            show_info("暂无保存的模板")


def render_element_selector(elements: List[Dict], location_mapping: Dict, selected_id: str):
    """渲染段落选择器"""
    for i, elem in enumerate(elements):
        elem_id = elem["element_id"]
        text = elem["text"]
        is_selected = selected_id == elem_id
        
        mapped_vars = [
            k for k, v in location_mapping.items() 
            if v.get("element_id") == elem_id
        ]
        
        c1, c2 = st.columns([0.5, 9.5])
        
        with c1:
            btn_label = "📍" if is_selected else str(i+1)
            if st.button(btn_label, key=f"sel_{elem_id}"):
                st.session_state.selected_element_id = elem_id
                st.rerun()
        
        with c2:
            display_text = text[:80] + "..." if len(text) > 80 else text
            if is_selected:
                st.markdown(f"**{display_text}**")
            else:
                st.text(display_text)
            
            if mapped_vars:
                st.markdown(" ".join([f"`{v}`" for v in mapped_vars]))
        
        st.divider()


def render_mapping_config(elem: Dict, elem_id: str, location_mapping: Dict):
    """渲染映射配置面板"""
    st.markdown(f"**选中段落:**")
    show_info(elem["text"][:100] + ("..." if len(elem["text"]) > 100 else ""))
    
    # 自定义映射
    st.markdown("#### ✏️ 自定义映射")
    st.caption("手动输入要替换的文本和变量名")
    
    c1, c2 = st.columns(2)
    with c1:
        custom_text = st.text_input("原文本", key="custom_text", placeholder="如：陈长")
    with c2:
        custom_var = st.text_input("变量名", key="custom_var", placeholder="如：姓名")
    
    if st.button("➕ 添加自定义映射", type="primary", use_container_width=True):
        if custom_text and custom_var:
            pos = elem["text"].find(custom_text)
            if pos >= 0:
                location_mapping[custom_var] = {
                    "element_id": elem_id,
                    "start": pos,
                    "end": pos + len(custom_text),
                    "length": len(custom_text),
                    "original_text": custom_text
                }
                st.session_state.location_mapping = location_mapping
                show_success(f"已添加: {custom_var} = {custom_text}")
                st.rerun()
            else:
                show_error(f"未找到文本: {custom_text}")
    
    st.divider()
    
    # 智能检测
    st.markdown("#### 🔍 智能检测")
    st.caption("自动识别可替换内容")
    
    candidates = extract_candidates(elem["text"])
    
    if candidates:
        for cand in candidates:
            c1, c2, c3 = st.columns([3, 2, 1])
            with c1:
                st.markdown(f"`{cand['text']}` ({cand['type']})")
            with c2:
                var_input = st.text_input(
                    "变量名", 
                    key=f"det_var_{cand['start']}",
                    placeholder="输入变量名",
                    label_visibility="collapsed"
                )
            with c3:
                if st.button("使用", key=f"det_add_{cand['start']}"):
                    if var_input:
                        location_mapping[var_input] = {
                            "element_id": elem_id,
                            "start": cand["start"],
                            "end": cand["end"],
                            "length": cand["end"] - cand["start"],
                            "original_text": cand["text"]
                        }
                        st.session_state.location_mapping = location_mapping
                        st.rerun()
    else:
        show_info("未检测到可替换内容")


def render_mapping_list(location_mapping: Dict):
    """渲染已配置映射列表"""
    st.divider()
    st.markdown("#### 📋 已配置映射")
    
    if location_mapping:
        for var_name, loc in location_mapping.items():
            c1, c2 = st.columns([4, 1])
            with c1:
                st.write(f"**{var_name}** = `{loc['original_text']}`")
            with c2:
                if st.button("🗑️", key=f"del_map_{var_name}"):
                    del st.session_state.location_mapping[var_name]
                    st.rerun()
    else:
        show_info("暂无映射配置")


def render_template_page():
    """渲染模板管理页面"""
    st.header("📄 步骤1: 模板管理")
    
    # 已保存模板
    render_saved_templates()
    
    st.divider()
    
    # 上传
    uploaded_file = st.file_uploader("📤 上传Word合同 (.docx)", type=["docx"])
    
    if uploaded_file:
        file_bytes = uploaded_file.getvalue()
        
        if st.session_state.uploaded_template_bytes != file_bytes:
            st.session_state.uploaded_template_bytes = file_bytes
            st.session_state.doc_elements = parse_doc_elements(file_bytes)
            st.session_state.location_mapping = {}
            st.session_state.template_name = Path(uploaded_file.name).stem
            st.session_state.selected_element_id = None
        
        st.markdown(PAGE_STYLE, unsafe_allow_html=True)
        
        # 双列布局
        col_preview, col_config = st.columns([3, 2])
        
        with col_preview:
            st.subheader("📄 合同预览")
            st.caption("点击段落编号选择，在右侧配置映射")
            
            with st.container():
                render_element_selector(
                    st.session_state.doc_elements,
                    st.session_state.location_mapping,
                    st.session_state.selected_element_id
                )
        
        with col_config:
            st.subheader("🏷️ 配置映射")
            
            if st.session_state.selected_element_id:
                elem_id = st.session_state.selected_element_id
                elem = next((e for e in st.session_state.doc_elements if e["element_id"] == elem_id), None)
                
                if elem:
                    render_mapping_config(elem, elem_id, st.session_state.location_mapping)
            else:
                show_info("👈 请在左侧点击段落编号选择")
            
            render_mapping_list(st.session_state.location_mapping)
        
        # 保存模板
        st.divider()
        st.subheader("💾 保存模板")
        
        c1, c2 = st.columns([2, 1])
        with c1:
            st.session_state.template_name = st.text_input("模板名称 *", value=st.session_state.template_name)
        with c2:
            st.session_state.description = st.text_input("描述", value=st.session_state.description)
        
        if st.button("💾 保存模板", type="primary", use_container_width=True):
            if not st.session_state.template_name:
                show_error("请输入模板名称")
            elif not st.session_state.location_mapping:
                show_error("请至少添加一个映射")
            else:
                try:
                    config = template_service.create_location_template(
                        template_name=st.session_state.template_name,
                        original_filename=uploaded_file.name,
                        docx_bytes=file_bytes,
                        location_mapping=st.session_state.location_mapping,
                        description=st.session_state.description
                    )
                    show_success(f"保存成功！ID: {config.template_id}")
                except Exception as e:
                    show_error(f"保存失败: {e}")
