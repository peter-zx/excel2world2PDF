"""
Word文档处理服务
重点：保留原始格式进行替换
"""
from io import BytesIO
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import re


class WordService:
    """Word文档处理服务"""
    
    def replace_preserving_format(
        self,
        file_bytes: bytes,
        mapping: Dict[str, str],  # {变量名: 新值}
        location_mapping: Optional[Dict[str, Dict]] = None,
        text_mapping: Optional[Dict[str, str]] = None
    ) -> bytes:
        """
        替换内容并保留格式
        """
        doc = Document(BytesIO(file_bytes))
        
        # 构建元素映射
        element_map = {}
        element_texts = {}
        
        for para_idx, para in enumerate(doc.paragraphs):
            elem_id = f"para_{para_idx}"
            element_map[elem_id] = para
            element_texts[elem_id] = para.text
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    elem_id = f"cell_{table_idx}_{row_idx}_{cell_idx}"
                    element_map[elem_id] = cell
                    element_texts[elem_id] = cell.text
        
        # 构建位置映射
        if location_mapping:
            final_mapping = location_mapping
        elif text_mapping:
            final_mapping = self._build_location_mapping(element_texts, text_mapping)
        else:
            final_mapping = {}
        
        # 执行替换
        for var_name, new_value in mapping.items():
            if var_name not in final_mapping:
                continue
            
            loc = final_mapping[var_name]
            elem_id = loc["element_id"]
            
            if elem_id not in element_map:
                continue
            
            element = element_map[elem_id]
            
            if elem_id.startswith("para_"):
                self._replace_in_paragraph_preserve_format(element, loc, new_value)
            elif elem_id.startswith("cell_"):
                self._replace_in_cell_preserve_format(element, loc, new_value)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
    
    def _replace_in_paragraph_preserve_format(self, paragraph, location: Dict, new_text: str):
        """
        在段落中替换文本，保留格式
        
        关键：找到目标位置所在的run，保留该run的格式属性
        """
        start = location["start"]
        end = location["end"]
        original_text = location.get("original_text", "")
        
        # 验证原文本是否在预期位置
        full_text = paragraph.text
        if full_text[start:end] != original_text:
            # 位置不匹配，尝试重新查找
            actual_pos = full_text.find(original_text)
            if actual_pos >= 0:
                start = actual_pos
                end = actual_pos + len(original_text)
            else:
                return
        
        # 找到目标run
        runs = paragraph.runs
        current_pos = 0
        target_runs = []  # 存储涉及的runs及其内部位置
        
        for run in runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            # 检查这个run是否与目标范围重叠
            if run_start < end and run_end > start:
                # 计算在run内的位置
                in_run_start = max(0, start - run_start)
                in_run_end = min(len(run.text), end - run_start)
                
                target_runs.append({
                    "run": run,
                    "run_start": run_start,
                    "run_end": run_end,
                    "in_start": in_run_start,
                    "in_end": in_run_end,
                    "is_first": run_start <= start,
                    "is_last": run_end >= end
                })
            
            current_pos = run_end
        
        if not target_runs:
            return
        
        # 执行替换
        if len(target_runs) == 1:
            # 最简单情况：目标文本在一个run内
            run_info = target_runs[0]
            run = run_info["run"]
            old_text = run.text
            run.text = old_text[:run_info["in_start"]] + new_text + old_text[run_info["in_end"]:]
        
        else:
            # 跨多个run的情况
            first_run = target_runs[0]
            last_run = target_runs[-1]
            
            # 在第一个run中放入替换文本
            run = first_run["run"]
            old_text = run.text
            if first_run["is_first"]:
                run.text = old_text[:first_run["in_start"]] + new_text
            else:
                run.text = new_text
            
            # 清空中间的runs
            for run_info in target_runs[1:-1]:
                run_info["run"].text = ""
            
            # 处理最后一个run
            if len(target_runs) > 1:
                run = last_run["run"]
                old_text = run.text
                run.text = old_text[last_run["in_end"]:]
    
    def _replace_in_cell_preserve_format(self, cell, location: Dict, new_text: str):
        """在表格单元格中替换，保留格式"""
        if cell.paragraphs:
            self._replace_in_paragraph_preserve_format(cell.paragraphs[0], location, new_text)
    
    def _build_location_mapping(self, element_texts: Dict[str, str], text_mapping: Dict[str, str]) -> Dict[str, Dict]:
        """从文本映射构建位置映射"""
        location_mapping = {}
        
        for var_name, original_text in text_mapping.items():
            if not original_text:
                continue
            
            for elem_id, text in element_texts.items():
                pos = text.find(original_text)
                if pos >= 0:
                    location_mapping[var_name] = {
                        "element_id": elem_id,
                        "start": pos,
                        "end": pos + len(original_text),
                        "length": len(original_text),
                        "original_text": original_text
                    }
                    break
        
        return location_mapping
    
    def batch_generate_by_location(
        self,
        template_bytes: bytes,
        data_list: List[Dict[str, str]],
        location_mapping: Dict[str, Dict]
    ) -> List[Tuple[str, bytes]]:
        """批量生成（位置映射模式）"""
        results = []
        
        for idx, data in enumerate(data_list):
            try:
                doc_bytes = self.replace_preserving_format(
                    template_bytes,
                    data,
                    location_mapping=location_mapping
                )
                
                name = data.get("姓名", data.get("name", f"合同_{idx+1}"))
                safe_name = "".join(
                    c for c in str(name) 
                    if c.isalnum() or c in (' ', '-', '_') or '\u4e00' <= c <= '\u9fff'
                )
                filename = f"{safe_name}_合同.docx"
                
                results.append((filename, doc_bytes))
                
            except Exception as e:
                print(f"Error generating contract {idx+1}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        return results
    
    def batch_generate_by_text(
        self,
        template_bytes: bytes,
        data_list: List[Dict[str, str]],
        text_mapping: Dict[str, str]
    ) -> List[Tuple[str, bytes]]:
        """批量生成（文本映射模式）"""
        results = []
        
        for idx, data in enumerate(data_list):
            try:
                doc_bytes = self.replace_preserving_format(
                    template_bytes,
                    data,
                    text_mapping=text_mapping
                )
                
                name = data.get("姓名", data.get("name", f"合同_{idx+1}"))
                safe_name = "".join(
                    c for c in str(name) 
                    if c.isalnum() or c in (' ', '-', '_') or '\u4e00' <= c <= '\u9fff'
                )
                filename = f"{safe_name}_合同.docx"
                
                results.append((filename, doc_bytes))
                
            except Exception as e:
                print(f"Error generating contract {idx+1}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        return results


# 单例
word_service = WordService()
